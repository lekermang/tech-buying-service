import io
import json
import os
import psycopg2
import requests

HEADERS = {'Access-Control-Allow-Origin': '*'}


STATUS_LABELS = {
    'new': '🆕 Новая',
    'in_progress': '🔧 В работе',
    'waiting_parts': '⏳ Ожидание запчастей',
    'ready': '✅ Готово к выдаче',
    'done': '✔️ Выдано',
    'cancelled': '❌ Отменено',
}


def send_tg_all(token: str, main_chat_id: str, conn, message: str):
    """Отправить уведомление в основной чат и всем активным получателям из БД"""
    tg_url = f'https://api.telegram.org/bot{token}'
    recipients = [main_chat_id]
    try:
        cur2 = conn.cursor()
        cur2.execute(
            f"SELECT telegram_chat_id FROM {SCHEMA}.notification_recipients WHERE is_active = true AND notify_repair = true"
        )
        rows = cur2.fetchall()
        cur2.close()
        for row in rows:
            cid = row[0]
            if cid and cid not in recipients:
                recipients.append(cid)
    except Exception:
        pass
    pluxan = os.environ.get('PLUXAN4IK_CHAT_ID', '')
    if pluxan and pluxan not in recipients:
        recipients.append(pluxan)
    for cid in recipients:
        try:
            requests.post(
                f'{tg_url}/sendMessage',
                json={'chat_id': cid, 'text': message, 'parse_mode': 'Markdown'},
                timeout=10
            )
        except Exception:
            pass


def build_act_docx(order_id, name, phone, model, repair_type, price_str, comment) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import datetime

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2); section.bottom_margin = Cm(2)
        section.left_margin = Cm(3); section.right_margin = Cm(1.5)

    def add_para(text='', bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph(); p.alignment = align
        run = p.add_run(text); run.bold = bold
        run.font.size = Pt(size); run.font.name = 'Times New Roman'
        return p

    def add_field(label, value):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p.add_run(label); r1.bold = True; r1.font.size = Pt(12); r1.font.name = 'Times New Roman'
        r2 = p.add_run(value); r2.font.size = Pt(12); r2.font.name = 'Times New Roman'

    now = datetime.datetime.now().strftime('%d.%m.%Y %H:%M')
    add_para('АКТ ПРИЁМКИ НА РЕМОНТ ТЕХНИЧЕСКОГО ОБОРУДОВАНИЯ', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(f'№ {order_id}', size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(f'г. Калуга, {now}', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_field('Исполнитель: ', 'ИП Мамедов Адиль Мирза Оглы, ИНН 402810962699, г. Калуга, ул. Кирова, 21а')
    add_field('Заказчик: ', name); add_field('Телефон: ', phone)
    if model: add_field('Устройство: ', model)
    if repair_type: add_field('Вид работ: ', repair_type)
    add_field('Предварительная стоимость: ', price_str)
    if comment: add_field('Описание: ', comment)
    doc.add_paragraph()
    add_para('УСЛОВИЯ РЕМОНТА — КЛИЕНТ ОЗНАКОМЛЕН И СОГЛАСЕН:', bold=True, size=13)
    doc.add_paragraph()
    risks = [
        'После контакта с водой аппарат может полностью выйти из строя при любом виде ремонта. Мастерская не обязана восстанавливать устройство.',
        'Компонентная пайка — риск безвозвратного повреждения платы. Если телефон перестал включаться в процессе — работа оплачивается.',
        'При снятии дисплея возможно повреждение (полосы, артефакты). Замена за счёт клиента.',
        'Все данные (фото, контакты) могут быть безвозвратно утеряны. Мастерская не восстанавливает данные.',
        'Гарантия на результат не предоставляется. В худшем случае клиент оплачивает диагностику и выполненные работы.',
    ]
    for i, risk in enumerate(risks, 1):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'{i}. {risk}'); run.font.size = Pt(12); run.font.name = 'Times New Roman'
    doc.add_paragraph()
    add_para('Подписывая акт, клиент подтверждает ознакомление с условиями и добровольно соглашается на ремонт.', size=11)
    doc.add_paragraph(); doc.add_paragraph()

    table = doc.add_table(rows=2, cols=3)
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border in ['top', 'bottom', 'left', 'right']:
                b = OxmlElement(f'w:{border}'); b.set(qn('w:val'), 'nil'); tcBorders.append(b)
            tcPr.append(tcBorders)
    for col_idx, label in [(0, 'Мастер (подпись / ФИО)'), (2, 'Клиент (подпись / ФИО)')]:
        cell = table.cell(0, col_idx); p = cell.paragraphs[0]
        run = p.add_run('_' * 30); run.font.size = Pt(12)
        cell2 = table.cell(1, col_idx); p2 = cell2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(label); run2.font.size = Pt(10); run2.font.name = 'Times New Roman'

    doc.add_paragraph()
    p_f = doc.add_paragraph(); p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_f.add_run('ИНН: 402810962699  ·  ОГРНИП: 307402814200032  ·  Р/с: 40802810422270001866\nКАЛУЖСКОЕ ОТДЕЛЕНИЕ N8608 ПАО СБЕРБАНК  ·  БИК: 042908612')
    r.font.size = Pt(9); r.font.name = 'Times New Roman'

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


def send_tg_document(token: str, chat_id, doc_bytes: bytes, filename: str, caption: str = ''):
    try:
        requests.post(
            f'https://api.telegram.org/bot{token}/sendDocument',
            data={'chat_id': chat_id, 'caption': caption},
            files={'document': (filename, doc_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')},
            timeout=30,
        )
    except Exception:
        pass


def send_sms(phone: str, message: str):
    """Отправить SMS через sms.ru"""
    api_id = os.environ.get('SMSRU_API_ID', '')
    if not api_id:
        return
    clean_phone = ''.join(c for c in phone if c.isdigit() or c == '+')
    try:
        requests.get(
            'https://sms.ru/sms/send',
            params={'api_id': api_id, 'to': clean_phone, 'msg': message, 'json': 1, 'from': 'Skypka24'},
            timeout=10
        )
    except Exception:
        pass
SCHEMA = 't_p31606708_tech_buying_service'
ADMIN_TOKEN = os.environ.get('ADMIN_TOKEN', 'Mark2015N')
ADMIN_TOKEN_ALT = 'Mark2015N'

VALID_STATUSES = ['new', 'in_progress', 'waiting_parts', 'ready', 'done', 'cancelled']
ALLOW_HEADERS = 'Content-Type, X-Admin-Token, X-Employee-Token'


def auth(event: dict) -> bool:
    headers = {k.lower(): v for k, v in (event.get('headers') or {}).items()}
    admin_token = headers.get('x-admin-token', '')
    if admin_token and (admin_token == ADMIN_TOKEN or admin_token == ADMIN_TOKEN_ALT):
        return True
    emp_token = headers.get('x-employee-token', '')
    if emp_token:
        conn = psycopg2.connect(os.environ['DATABASE_URL'])
        cur = conn.cursor()
        cur.execute(
            f"SELECT id FROM {SCHEMA}.employees WHERE auth_token='{emp_token}' AND token_expires_at>NOW() AND is_active=true"
        )
        row = cur.fetchone()
        cur.close(); conn.close()
        return row is not None
    return False


def is_owner(event: dict) -> bool:
    headers = {k.lower(): v for k, v in (event.get('headers') or {}).items()}
    admin_token = headers.get('x-admin-token', '')
    if admin_token and (admin_token == ADMIN_TOKEN or admin_token == ADMIN_TOKEN_ALT):
        return True
    emp_token = headers.get('x-employee-token', '')
    if emp_token:
        conn = psycopg2.connect(os.environ['DATABASE_URL'])
        cur = conn.cursor()
        cur.execute(
            f"SELECT role FROM {SCHEMA}.employees WHERE auth_token='{emp_token}' AND token_expires_at>NOW() AND is_active=true"
        )
        row = cur.fetchone()
        cur.close(); conn.close()
        return row is not None and row[0] == 'owner'
    return False


def handler(event: dict, context) -> dict:
    """Управление заявками на ремонт: список, создание, смена статуса, аналитика по периодам, доход мастера"""

    if event.get('httpMethod') == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {**HEADERS, 'Access-Control-Allow-Methods': 'GET, POST, OPTIONS', 'Access-Control-Allow-Headers': ALLOW_HEADERS},
            'body': '',
        }

    # Публичные actions (без токена)
    raw_body_pub = event.get('body') or '{}'
    body_pub = json.loads(raw_body_pub) if isinstance(raw_body_pub, str) else (raw_body_pub or {})
    if event.get('httpMethod') == 'POST' and body_pub.get('action') == 'client_register':
        conn = psycopg2.connect(os.environ['DATABASE_URL'])
        cur = conn.cursor()
        full_name = str(body_pub.get('full_name', '')).strip().replace("'", "''")
        phone = ''.join(c for c in str(body_pub.get('phone', '')) if c.isdigit())
        birth_date = str(body_pub.get('birth_date', '')).strip()
        ref_phone = ''.join(c for c in str(body_pub.get('referrer_phone', '')) if c.isdigit())
        if not full_name or len(phone) < 10:
            cur.close(); conn.close()
            return {'statusCode': 400, 'headers': HEADERS, 'body': json.dumps({'error': 'ФИО и телефон обязательны'}, ensure_ascii=False)}
        referrer_id = 'NULL'
        if ref_phone and len(ref_phone) >= 10:
            ref_suffix = ref_phone[-10:]
            cur.execute(f"SELECT id FROM {SCHEMA}.repair_clients WHERE RIGHT(REGEXP_REPLACE(phone, '[^0-9]', '', 'g'), 10)='{ref_suffix}'")
            ref_row = cur.fetchone()
            if ref_row:
                referrer_id = str(ref_row[0])
                cur.execute(f"UPDATE {SCHEMA}.repair_clients SET discount_pct=GREATEST(discount_pct, 5) WHERE id={ref_row[0]}")
        bd_val = f"'{birth_date}'" if birth_date else 'NULL'
        cur.execute(f"""
            INSERT INTO {SCHEMA}.repair_clients (full_name, phone, birth_date, referrer_id, discount_pct)
            VALUES ('{full_name}', '+{phone}', {bd_val}, {referrer_id}, 3)
            ON CONFLICT (phone) DO UPDATE SET full_name=EXCLUDED.full_name, birth_date=EXCLUDED.birth_date
            RETURNING id, discount_pct
        """)
        row = cur.fetchone()
        conn.commit(); cur.close(); conn.close()
        return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'ok': True, 'client_id': row[0], 'discount_pct': row[1]}, ensure_ascii=False)}

    if not auth(event):
        return {'statusCode': 401, 'headers': HEADERS, 'body': json.dumps({'error': 'Unauthorized'}, ensure_ascii=False)}

    method = event.get('httpMethod', 'GET')
    params = event.get('queryStringParameters') or {}

    conn = psycopg2.connect(os.environ['DATABASE_URL'])
    cur = conn.cursor()

    if method == 'GET':
        action = params.get('action', '')

        # Аналитика за период (day/week/month)
        if action == 'analytics':
            period = params.get('period', 'month')
            if period == 'day':
                interval = "1 day"
            elif period == 'week':
                interval = "7 days"
            else:
                interval = "30 days"

            cur.execute(f"""
                SELECT
                    COUNT(*) as total,
                    COUNT(*) FILTER (WHERE status = 'done') as done,
                    COUNT(*) FILTER (WHERE status = 'cancelled') as cancelled,
                    COUNT(*) FILTER (WHERE status = 'ready') as ready,
                    COUNT(*) FILTER (WHERE status = 'in_progress') as in_progress,
                    COUNT(*) FILTER (WHERE status = 'waiting_parts') as waiting_parts,
                    COUNT(*) FILTER (WHERE status = 'new') as new_count,
                    COALESCE(SUM(repair_amount) FILTER (WHERE status = 'done'), 0) as revenue,
                    COALESCE(SUM(purchase_amount) FILTER (WHERE status = 'done'), 0) as costs,
                    COALESCE(SUM(master_income) FILTER (WHERE status = 'done'), 0) as master_total
                FROM {SCHEMA}.repair_orders
                WHERE created_at >= NOW() - INTERVAL '{interval}'
            """)
            row = cur.fetchone()

            revenue = int(row[7]) if row[7] else 0
            costs = int(row[8]) if row[8] else 0
            master_total = int(row[9]) if row[9] else 0
            profit = revenue - costs

            # Динамика по дням
            cur.execute(f"""
                SELECT
                    DATE(created_at + INTERVAL '3 hours') as day,
                    COUNT(*) as total,
                    COUNT(*) FILTER (WHERE status = 'done') as done,
                    COALESCE(SUM(repair_amount) FILTER (WHERE status = 'done'), 0) as revenue,
                    COALESCE(SUM(purchase_amount) FILTER (WHERE status = 'done'), 0) as costs
                FROM {SCHEMA}.repair_orders
                WHERE created_at >= NOW() - INTERVAL '{interval}'
                GROUP BY day
                ORDER BY day ASC
            """)
            daily_rows = cur.fetchall()
            daily = [
                {
                    'day': str(r[0]),
                    'total': r[1],
                    'done': r[2],
                    'revenue': int(r[3]),
                    'costs': int(r[4]),
                    'profit': int(r[3]) - int(r[4]),
                }
                for r in daily_rows
            ]

            cur.close(); conn.close()
            return {
                'statusCode': 200,
                'headers': HEADERS,
                'body': json.dumps({
                    'period': period,
                    'total': row[0], 'done': row[1], 'cancelled': row[2],
                    'ready': row[3], 'in_progress': row[4], 'waiting_parts': row[5], 'new': row[6],
                    'revenue': revenue, 'costs': costs, 'profit': profit, 'master_total': master_total,
                    'daily': daily,
                }, ensure_ascii=False)
            }

        # Дневная статистика (30 дней)
        if action == 'daily_stats':
            cur.execute(f"""
                SELECT
                    DATE(created_at + INTERVAL '3 hours') as day,
                    COUNT(*) as total,
                    COUNT(*) FILTER (WHERE status = 'done') as done,
                    COUNT(*) FILTER (WHERE status = 'cancelled') as cancelled,
                    COALESCE(SUM(repair_amount) FILTER (WHERE status = 'done'), 0) as revenue,
                    COALESCE(SUM(purchase_amount) FILTER (WHERE status = 'done'), 0) as costs,
                    COALESCE(SUM(master_income) FILTER (WHERE status = 'done'), 0) as master_income
                FROM {SCHEMA}.repair_orders
                WHERE created_at >= NOW() - INTERVAL '30 days'
                GROUP BY day
                ORDER BY day DESC
            """)
            rows = cur.fetchall()
            cur.close(); conn.close()
            stats = [
                {
                    'day': str(r[0]), 'total': r[1], 'done': r[2],
                    'cancelled': r[3], 'revenue': int(r[4]), 'costs': int(r[5]),
                    'profit': int(r[4]) - int(r[5]),
                    'master_income': int(r[6]),
                }
                for r in rows
            ]
            return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'stats': stats}, ensure_ascii=False)}

        # Настройки системы
        if action == 'settings_get':
            cur.execute(f"SELECT key, value, description, updated_at FROM {SCHEMA}.settings ORDER BY key")
            rows = cur.fetchall()
            cur.close(); conn.close()
            settings = [{'key': r[0], 'value': r[1], 'description': r[2], 'updated_at': r[3].isoformat() if r[3] else None} for r in rows]
            return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'settings': settings}, ensure_ascii=False)}

        # Цены работ + наценка на детали + доп. работы
        if action == 'labor_prices_get':
            cur.execute(f"SELECT part_type, label, price FROM {SCHEMA}.repair_labor_prices ORDER BY part_type")
            rows = cur.fetchall()
            cur.execute(f"SELECT value FROM {SCHEMA}.settings WHERE key='parts_markup'")
            mrow = cur.fetchone()
            markup = int(mrow[0]) if mrow else 0
            cur.execute(f"SELECT id, label, price, is_active, sort_order FROM {SCHEMA}.repair_extra_works ORDER BY sort_order, id")
            erows = cur.fetchall()
            extra = [{'id': r[0], 'label': r[1], 'price': r[2], 'is_active': r[3], 'sort_order': r[4]} for r in erows]
            cur.close(); conn.close()
            return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps(
                {'prices': [{'part_type': r[0], 'label': r[1], 'price': r[2]} for r in rows],
                 'parts_markup': markup, 'extra_works': extra},
                ensure_ascii=False
            )}

        # Список получателей уведомлений
        if action == 'recipients':
            cur.execute(
                f"SELECT id, name, telegram_chat_id, is_active, notify_repair, created_at FROM {SCHEMA}.notification_recipients ORDER BY created_at"
            )
            rows = cur.fetchall()
            cur.close(); conn.close()
            recipients = [
                {'id': r[0], 'name': r[1], 'telegram_chat_id': r[2], 'is_active': r[3], 'notify_repair': r[4], 'created_at': r[5].isoformat() if r[5] else None}
                for r in rows
            ]
            return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'recipients': recipients}, ensure_ascii=False)}

        # SMS контакты для рассылки
        if action == 'sms_contacts':
            group = params.get('group', 'all')

            def fmt_phone(raw: str) -> str:
                digits = ''.join(c for c in (raw or '') if c.isdigit())
                if len(digits) == 11 and digits.startswith('8'):
                    digits = '7' + digits[1:]
                if len(digits) == 10:
                    digits = '7' + digits
                return ('+' + digits) if len(digits) == 11 else ''

            contacts = []
            seen_phones = set()
            if group in ('all', 'registered'):
                cur.execute(f"SELECT id, full_name, phone FROM {SCHEMA}.clients WHERE (client_group IS NULL OR client_group != 'wh') ORDER BY registered_at DESC")
                for r in cur.fetchall():
                    p = fmt_phone(r[2] or '')
                    if p and p not in seen_phones:
                        contacts.append({'id': f'c_{r[0]}', 'full_name': r[1] or '', 'phone': p, 'source': 'registered'})
                        seen_phones.add(p)
            if group in ('all', 'repair'):
                cur.execute(
                    "SELECT id, name, phone FROM " + SCHEMA + ".repair_orders"
                    " WHERE status NOT IN ('cancelled')"
                    " AND phone IS NOT NULL AND phone != '' AND LENGTH(phone) >= 11"
                    " ORDER BY id DESC"
                )
                for r in cur.fetchall():
                    p = fmt_phone((r[2] or '').strip())
                    if p and p not in seen_phones:
                        contacts.append({'id': f'r_{r[0]}', 'full_name': r[1] or '', 'phone': p, 'source': 'repair'})
                        seen_phones.add(p)
            if group in ('all', 'wh'):
                cur.execute(f"SELECT id, full_name, phone FROM {SCHEMA}.clients WHERE client_group = 'wh' ORDER BY full_name")
                for r in cur.fetchall():
                    p = fmt_phone(r[2] or '')
                    if p and p not in seen_phones:
                        contacts.append({'id': f'c_{r[0]}', 'full_name': r[1] or '', 'phone': p, 'source': 'wh'})
                        seen_phones.add(p)
            cur.close(); conn.close()
            print(f"[sms_contacts] group={group} total={len(contacts)}", flush=True)
            return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'contacts': contacts, 'total': len(contacts)}, ensure_ascii=False)}

        # Список заявок
        status_filter = params.get('status', '')
        search = params.get('search', '')
        if status_filter and search:
            cur.execute(
                f"SELECT id, name, phone, model, repair_type, price, status, admin_note, created_at, comment, purchase_amount, repair_amount, completed_at, master_income, parts_name FROM {SCHEMA}.repair_orders WHERE status = '{status_filter}' AND (name ILIKE '%{search}%' OR phone ILIKE '%{search}%' OR model ILIKE '%{search}%') ORDER BY created_at DESC LIMIT 200"
            )
        elif status_filter:
            cur.execute(
                f"SELECT id, name, phone, model, repair_type, price, status, admin_note, created_at, comment, purchase_amount, repair_amount, completed_at, master_income, parts_name FROM {SCHEMA}.repair_orders WHERE status = '{status_filter}' ORDER BY created_at DESC LIMIT 200"
            )
        elif search:
            cur.execute(
                f"SELECT id, name, phone, model, repair_type, price, status, admin_note, created_at, comment, purchase_amount, repair_amount, completed_at, master_income, parts_name FROM {SCHEMA}.repair_orders WHERE name ILIKE '%{search}%' OR phone ILIKE '%{search}%' OR model ILIKE '%{search}%' ORDER BY created_at DESC LIMIT 200"
            )
        else:
            cur.execute(
                f"SELECT id, name, phone, model, repair_type, price, status, admin_note, created_at, comment, purchase_amount, repair_amount, completed_at, master_income, parts_name FROM {SCHEMA}.repair_orders ORDER BY created_at DESC LIMIT 200"
            )
        rows = cur.fetchall()
        cur.close(); conn.close()

        orders = [
            {
                'id': r[0], 'name': r[1], 'phone': r[2], 'model': r[3],
                'repair_type': r[4], 'price': r[5], 'status': r[6],
                'admin_note': r[7], 'created_at': r[8].isoformat() if r[8] else None,
                'comment': r[9], 'purchase_amount': r[10], 'repair_amount': r[11],
                'completed_at': r[12].isoformat() if r[12] else None,
                'master_income': r[13], 'parts_name': r[14],
            }
            for r in rows
        ]
        return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'orders': orders}, ensure_ascii=False)}

    if method == 'POST':
        raw_body = event.get('body') or '{}'
        body = json.loads(raw_body) if isinstance(raw_body, str) else (raw_body or {})
        action = body.get('action', 'update_status')

        # Сохранение цен работ + наценки + доп. работ + мгновенный пересчёт repair_parts
        if action == 'labor_prices_set':
            prices = body.get('prices', [])
            for item in prices:
                pt = str(item.get('part_type', '')).strip()
                price_val = int(item.get('price', 0))
                if pt:
                    cur.execute(f"UPDATE {SCHEMA}.repair_labor_prices SET price={price_val}, updated_at=NOW() WHERE part_type='{pt}'")
                    # Мгновенно обновляем labor_cost во всех запчастях этого типа
                    cur.execute(f"UPDATE {SCHEMA}.repair_parts SET labor_cost={price_val} WHERE part_type='{pt}'")
            if 'parts_markup' in body:
                markup_val = int(body.get('parts_markup', 0))
                cur.execute(f"""
                    INSERT INTO {SCHEMA}.settings (key, value, description)
                    VALUES ('parts_markup', '{markup_val}', 'Наценка на запчасти для ремонта (руб)')
                    ON CONFLICT (key) DO UPDATE SET value='{markup_val}', updated_at=NOW()
                """)
            # Сохраняем доп. работы
            extra = body.get('extra_works', [])
            for ew in extra:
                eid = ew.get('id')
                elabel = str(ew.get('label', '')).replace("'", "''").strip()
                eprice = int(ew.get('price', 0))
                eactive = bool(ew.get('is_active', True))
                esort = int(ew.get('sort_order', 0))
                if eid:
                    cur.execute(f"""
                        UPDATE {SCHEMA}.repair_extra_works
                        SET label='{elabel}', price={eprice}, is_active={eactive}, sort_order={esort}
                        WHERE id={eid}
                    """)
                elif elabel:
                    cur.execute(f"""
                        INSERT INTO {SCHEMA}.repair_extra_works (label, price, is_active, sort_order)
                        VALUES ('{elabel}', {eprice}, {eactive}, {esort})
                    """)
            conn.commit()
            cur.close(); conn.close()
            return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'ok': True}, ensure_ascii=False)}

        # Удаление доп. работы
        if action == 'extra_work_delete':
            eid = int(body.get('id', 0))
            if eid:
                cur.execute(f"DELETE FROM {SCHEMA}.repair_extra_works WHERE id={eid}")
                conn.commit()
            cur.close(); conn.close()
            return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'ok': True}, ensure_ascii=False)}

        # Регистрация клиента (публичный action — без токена)
        if action == 'client_register':
            full_name = str(body.get('full_name', '')).strip().replace("'", "''")
            phone = ''.join(c for c in str(body.get('phone', '')) if c.isdigit())
            birth_date = str(body.get('birth_date', '')).strip()
            ref_phone = ''.join(c for c in str(body.get('referrer_phone', '')) if c.isdigit())
            if not full_name or len(phone) < 10:
                cur.close(); conn.close()
                return {'statusCode': 400, 'headers': HEADERS, 'body': json.dumps({'error': 'ФИО и телефон обязательны'}, ensure_ascii=False)}
            # Ищем реферера
            referrer_id = 'NULL'
            if ref_phone and len(ref_phone) >= 10:
                ref_suffix = ref_phone[-10:]
                cur.execute(f"SELECT id FROM {SCHEMA}.repair_clients WHERE RIGHT(REGEXP_REPLACE(phone, '[^0-9]', '', 'g'), 10)='{ref_suffix}'")
                ref_row = cur.fetchone()
                if ref_row:
                    referrer_id = str(ref_row[0])
                    # Обновляем скидку реферера до 5%
                    cur.execute(f"UPDATE {SCHEMA}.repair_clients SET discount_pct=GREATEST(discount_pct, 5) WHERE id={ref_row[0]}")
            bd_val = f"'{birth_date}'" if birth_date else 'NULL'
            cur.execute(f"""
                INSERT INTO {SCHEMA}.repair_clients (full_name, phone, birth_date, referrer_id, discount_pct)
                VALUES ('{full_name}', '+{phone}', {bd_val}, {referrer_id}, 3)
                ON CONFLICT (phone) DO UPDATE SET full_name=EXCLUDED.full_name, birth_date=EXCLUDED.birth_date
                RETURNING id, discount_pct
            """)
            row = cur.fetchone()
            conn.commit()
            cur.close(); conn.close()
            return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'ok': True, 'client_id': row[0], 'discount_pct': row[1]}, ensure_ascii=False)}

        # Тест SMS
        if action == 'sms_test':
            phone = (body.get('phone') or '').strip()
            message = (body.get('message') or 'Тестовое SMS от Скупка24 — всё работает!').strip()
            if not phone:
                cur.close(); conn.close()
                return {'statusCode': 400, 'headers': HEADERS, 'body': json.dumps({'error': 'Укажите phone'}, ensure_ascii=False)}
            api_id = os.environ.get('SMSRU_API_ID', '')
            if not api_id:
                cur.close(); conn.close()
                return {'statusCode': 500, 'headers': HEADERS, 'body': json.dumps({'error': 'SMSRU_API_ID не задан'}, ensure_ascii=False)}
            resp = requests.get(
                'https://sms.ru/sms/send',
                params={'api_id': api_id, 'to': phone, 'msg': message, 'json': 1},
                timeout=15
            )
            cur.close(); conn.close()
            return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'smsru_response': resp.json()}, ensure_ascii=False)}

        # SMS рассылка
        if action == 'sms_blast':
          try:
            print(f"[sms_blast] start group={body.get('group')} msg_len={len(body.get('message',''))}", flush=True)
            message = (body.get('message') or '').strip()
            group = body.get('group', 'all')
            if not message:
                cur.close(); conn.close()
                return {'statusCode': 400, 'headers': HEADERS, 'body': json.dumps({'error': 'Текст сообщения обязателен'}, ensure_ascii=False)}
            if len(message) > 480:
                cur.close(); conn.close()
                return {'statusCode': 400, 'headers': HEADERS, 'body': json.dumps({'error': 'Сообщение слишком длинное (макс 480 символов)'}, ensure_ascii=False)}
            api_id = os.environ.get('SMSRU_API_ID', '')
            if not api_id:
                cur.close(); conn.close()
                return {'statusCode': 500, 'headers': HEADERS, 'body': json.dumps({'error': 'SMS сервис не настроен'}, ensure_ascii=False)}

            def fmt_p(raw: str) -> str:
                digits = ''.join(c for c in (raw or '') if c.isdigit())
                if len(digits) == 11 and digits.startswith('8'):
                    digits = '7' + digits[1:]
                if len(digits) == 10:
                    digits = '7' + digits
                return ('+' + digits) if len(digits) == 11 else ''

            phones = set()
            if group in ('all', 'registered'):
                cur.execute(f"SELECT phone FROM {SCHEMA}.clients WHERE (client_group IS NULL OR client_group != 'wh')")
                for r in cur.fetchall():
                    p = fmt_p(r[0] or '')
                    if p: phones.add(p)
            if group in ('all', 'repair'):
                cur.execute(
                    "SELECT DISTINCT phone FROM " + SCHEMA + ".repair_orders"
                    " WHERE status NOT IN ('cancelled') AND phone IS NOT NULL AND LENGTH(phone) >= 11"
                )
                for r in cur.fetchall():
                    p = fmt_p((r[0] or '').strip())
                    if p: phones.add(p)
            if group in ('all', 'wh'):
                cur.execute(f"SELECT phone FROM {SCHEMA}.clients WHERE client_group = 'wh'")
                for r in cur.fetchall():
                    p = fmt_p(r[0] or '')
                    if p: phones.add(p)
            cur.close(); conn.close()
            all_phones = list(phones)
            sent = 0; failed = 0
            # sms.ru поддерживает до 100 номеров в одном запросе через запятую
            chunk_size = 100
            for i in range(0, len(all_phones), chunk_size):
                chunk = all_phones[i:i + chunk_size]
                try:
                    resp = requests.get('https://sms.ru/sms/send',
                        params={'api_id': api_id, 'to': ','.join(chunk), 'msg': message, 'json': 1, 'from': 'Skypka24'},
                        timeout=20)
                    data_r = resp.json() if resp.status_code == 200 else {}
                    print(f"[sms_blast] chunk resp status={data_r.get('status')} status_code={resp.status_code} body={str(data_r)[:300]}", flush=True)
                    if data_r.get('status') == 'OK':
                        sms_items = data_r.get('sms', {})
                        for num, info in sms_items.items():
                            if isinstance(info, dict) and info.get('status') == 'OK':
                                sent += 1
                            else:
                                failed += 1
                    else:
                        failed += len(chunk)
                except Exception:
                    failed += len(chunk)
            print(f"[sms_blast] sent={sent} failed={failed} total={len(all_phones)}", flush=True)
            return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'ok': True, 'sent': sent, 'failed': failed, 'total': len(all_phones)}, ensure_ascii=False)}
          except Exception as e:
            import traceback
            print(f"[sms_blast] EXCEPTION: {e}\n{traceback.format_exc()}", flush=True)
            return {'statusCode': 500, 'headers': HEADERS, 'body': json.dumps({'error': f'Внутренняя ошибка: {e}'}, ensure_ascii=False)}

        # Импорт WH-контактов из WhatsApp чата (Яндекс Диск)
        if action == 'import_wh_contacts':
            import re
            yd_url = (body.get('url') or '').strip()
            if not yd_url:
                cur.close(); conn.close()
                return {'statusCode': 400, 'headers': HEADERS, 'body': json.dumps({'error': 'url required'}, ensure_ascii=False)}
            # Получить прямую ссылку через API Яндекс Диска
            api_resp = requests.get(
                'https://cloud-api.yandex.net/v1/disk/public/resources/download',
                params={'public_key': yd_url}, timeout=15
            )
            if api_resp.status_code != 200:
                cur.close(); conn.close()
                return {'statusCode': 400, 'headers': HEADERS, 'body': json.dumps({'error': 'Не удалось получить ссылку с Яндекс Диска'}, ensure_ascii=False)}
            download_href = api_resp.json().get('href', '')
            if not download_href:
                cur.close(); conn.close()
                return {'statusCode': 400, 'headers': HEADERS, 'body': json.dumps({'error': 'Нет ссылки для скачивания'}, ensure_ascii=False)}
            file_resp = requests.get(download_href, timeout=30)
            text = file_resp.content.decode('utf-8', errors='ignore')
            # Извлечь все уникальные российские мобильные номера
            raw_numbers = re.findall(r'\b(?:\+?7|8)(9\d{9})\b', text)
            unique_phones = list(dict.fromkeys(['7' + n for n in raw_numbers]))
            if not unique_phones:
                cur.close(); conn.close()
                return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'ok': True, 'imported': 0, 'skipped': 0}, ensure_ascii=False)}
            # Получить уже существующие WH-номера
            cur.execute(f"SELECT phone FROM {SCHEMA}.clients WHERE client_group = 'wh'")
            existing = set(r[0] for r in cur.fetchall())
            # Получить максимальный номер WH
            cur.execute(f"SELECT full_name FROM {SCHEMA}.clients WHERE client_group = 'wh' ORDER BY id DESC LIMIT 1")
            last_row = cur.fetchone()
            last_num = 0
            if last_row:
                m = re.search(r'(\d+)$', last_row[0] or '')
                if m:
                    last_num = int(m.group(1))
            imported = 0; skipped = 0
            for phone in unique_phones:
                normalized = phone
                if normalized in existing:
                    skipped += 1
                    continue
                last_num += 1
                name = f'WH{last_num}'
                cur.execute(
                    f"INSERT INTO {SCHEMA}.clients (full_name, phone, client_group) VALUES ('{name}', '{normalized}', 'wh') ON CONFLICT DO NOTHING"
                )
                existing.add(normalized)
                imported += 1
            conn.commit(); cur.close(); conn.close()
            print(f"[import_wh_contacts] imported={imported} skipped={skipped}", flush=True)
            return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'ok': True, 'imported': imported, 'skipped': skipped}, ensure_ascii=False)}

        # Сохранить настройку
        if action == 'settings_set':
            key = (body.get('key') or '').strip()
            value = str(body.get('value') or '').strip()
            if not key:
                cur.close(); conn.close()
                return {'statusCode': 400, 'headers': HEADERS, 'body': json.dumps({'error': 'key required'}, ensure_ascii=False)}
            cur.execute(f"UPDATE {SCHEMA}.settings SET value='{value}', updated_at=NOW() WHERE key='{key}' RETURNING key")
            row = cur.fetchone()
            if not row:
                conn.commit(); cur.close(); conn.close()
                return {'statusCode': 404, 'headers': HEADERS, 'body': json.dumps({'error': 'setting not found'}, ensure_ascii=False)}
            updated_count = 0
            if key == 'price_markup':
                new_markup = int(value)
                cur.execute(
                    f"UPDATE {SCHEMA}.catalog SET price = original_price + {new_markup}, updated_at = NOW() WHERE original_price IS NOT NULL AND is_active = true"
                )
                updated_count = cur.rowcount
            conn.commit(); cur.close(); conn.close()
            return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'ok': True, 'key': key, 'value': value, 'prices_updated': updated_count}, ensure_ascii=False)}

        # Добавить получателя уведомлений
        if action == 'add_recipient':
            if not is_owner(event):
                cur.close(); conn.close()
                return {'statusCode': 403, 'headers': HEADERS, 'body': json.dumps({'error': 'Только владелец'}, ensure_ascii=False)}
            name = (body.get('name') or '').strip()
            chat_id_val = (body.get('telegram_chat_id') or '').strip()
            if not name or not chat_id_val:
                cur.close(); conn.close()
                return {'statusCode': 400, 'headers': HEADERS, 'body': json.dumps({'error': 'Укажите имя и chat_id'}, ensure_ascii=False)}
            cur.execute(
                f"INSERT INTO {SCHEMA}.notification_recipients (name, telegram_chat_id) VALUES ('{name}', '{chat_id_val}') RETURNING id"
            )
            new_id = cur.fetchone()[0]
            conn.commit(); cur.close(); conn.close()
            return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'ok': True, 'id': new_id}, ensure_ascii=False)}

        # Удалить получателя уведомлений
        if action == 'delete_recipient':
            if not is_owner(event):
                cur.close(); conn.close()
                return {'statusCode': 403, 'headers': HEADERS, 'body': json.dumps({'error': 'Только владелец'}, ensure_ascii=False)}
            rid = int(body.get('id', 0))
            cur.execute(f"DELETE FROM {SCHEMA}.notification_recipients WHERE id = {rid}")
            conn.commit(); cur.close(); conn.close()
            return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'ok': True}, ensure_ascii=False)}

        # Переключить активность получателя
        if action == 'toggle_recipient':
            if not is_owner(event):
                cur.close(); conn.close()
                return {'statusCode': 403, 'headers': HEADERS, 'body': json.dumps({'error': 'Только владелец'}, ensure_ascii=False)}
            rid = int(body.get('id', 0))
            cur.execute(f"UPDATE {SCHEMA}.notification_recipients SET is_active = NOT is_active WHERE id = {rid} RETURNING is_active")
            row = cur.fetchone()
            conn.commit(); cur.close(); conn.close()
            return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'ok': True, 'is_active': row[0] if row else False}, ensure_ascii=False)}

        # Создать заявку
        if action == 'create':
            name = (body.get('name') or '').strip()
            phone = (body.get('phone') or '').strip()
            if not name or not phone:
                cur.close(); conn.close()
                return {'statusCode': 400, 'headers': HEADERS, 'body': json.dumps({'error': 'Имя и телефон обязательны'}, ensure_ascii=False)}
            model = (body.get('model') or '').strip()
            repair_type = (body.get('repair_type') or '').strip()
            price = body.get('price')
            comment = (body.get('comment') or '').strip()
            price_val = int(price) if price else 'NULL'
            model_val = f"'{model}'" if model else 'NULL'
            repair_type_val = f"'{repair_type}'" if repair_type else 'NULL'
            comment_val = f"'{comment}'" if comment else 'NULL'
            price_str = f"{int(price):,} ₽".replace(',', ' ') if price else 'не определена'
            suffix = ''.join(c for c in phone if c.isdigit())[-10:]
            cur.execute(f"SELECT chat_id FROM {SCHEMA}.tg_phone_map WHERE RIGHT(REGEXP_REPLACE(phone, '[^0-9]', '', 'g'), 10) = '{suffix}' LIMIT 1")
            prow = cur.fetchone()
            client_chat_id = prow[0] if prow else None
            cur.execute(
                f"INSERT INTO {SCHEMA}.repair_orders (name, phone, model, repair_type, price, comment, client_tg_chat_id) VALUES ('{name}', '{phone}', {model_val}, {repair_type_val}, {price_val}, {comment_val}, {'NULL' if not client_chat_id else client_chat_id}) RETURNING id"
            )
            new_id = cur.fetchone()[0]
            conn.commit()
            tg_token = os.environ.get('TELEGRAM_BOT_TOKEN', '')
            main_chat = os.environ.get('TELEGRAM_CHAT_ID', '')
            tg_msg = (
                f"🔧 *Заявка #{new_id} на ремонт — Скупка24*\n\n"
                f"👤 *Имя:* {name}\n📞 *Телефон:* {phone}\n"
                f"📱 *Модель:* {model or '—'}\n🛠 *Тип ремонта:* {repair_type or '—'}\n"
                f"💰 *Стоимость:* {price_str}"
                + (f"\n📝 *Комментарий:* {comment}" if comment else "")
                + f"\n\n🔑 ID заявки: `{new_id}`"
                + ("\n✅ Клиент в TG" if client_chat_id else "")
            )
            if tg_token and main_chat:
                send_tg_all(tg_token, main_chat, conn, tg_msg)
                docx_bytes = build_act_docx(new_id, name, phone, model, repair_type, price_str, comment)
                filename = f'Акт_приёмки_{new_id}_{name.replace(" ", "_")}.docx'
                recipients = [main_chat]
                try:
                    cur3 = conn.cursor()
                    cur3.execute(f"SELECT telegram_chat_id FROM {SCHEMA}.notification_recipients WHERE is_active=true AND notify_repair=true")
                    for row in cur3.fetchall():
                        if row[0] and row[0] not in recipients:
                            recipients.append(row[0])
                    cur3.close()
                except Exception:
                    pass
                pluxan = os.environ.get('PLUXAN4IK_CHAT_ID', '')
                if pluxan and pluxan not in recipients:
                    recipients.append(pluxan)
                for cid in recipients:
                    send_tg_document(tg_token, cid, docx_bytes, filename, caption=f'📋 Акт приёмки №{new_id}')
            send_sms('+79929990333', f'Заявка #{new_id} на ремонт. {name}, {phone}. {repair_type or model or ""}. Скупка24')
            cur.close(); conn.close()
            return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'ok': True, 'id': new_id}, ensure_ascii=False)}

        # Удалить заявку
        if action == 'delete':
            order_id = int(body.get('id', 0))
            cur.execute(f"DELETE FROM {SCHEMA}.repair_orders WHERE id = {order_id}")
            conn.commit(); cur.close(); conn.close()
            return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'ok': True}, ensure_ascii=False)}

        # Обновить статус / поля заявки
        order_id = int(body.get('id', 0))
        new_status = body.get('status', '')
        admin_note = body.get('admin_note')
        purchase_amount = body.get('purchase_amount')
        repair_amount = body.get('repair_amount')
        parts_name = body.get('parts_name')
        # Базовые поля заявки
        upd_name = body.get('name')
        upd_phone = body.get('phone')
        upd_model = body.get('model')
        upd_repair_type = body.get('repair_type')
        upd_price = body.get('price')
        upd_comment = body.get('comment')

        if new_status and new_status not in VALID_STATUSES:
            cur.close(); conn.close()
            return {'statusCode': 400, 'headers': HEADERS, 'body': json.dumps({'error': 'Неверный статус'}, ensure_ascii=False)}

        # При переводе в ready — обязательные поля
        if new_status == 'ready':
            if not purchase_amount and purchase_amount != 0:
                cur.close(); conn.close()
                return {'statusCode': 400, 'headers': HEADERS, 'body': json.dumps({'error': 'Укажите сумму закупки запчасти', 'field': 'purchase_amount'}, ensure_ascii=False)}
            if not repair_amount and repair_amount != 0:
                cur.close(); conn.close()
                return {'statusCode': 400, 'headers': HEADERS, 'body': json.dumps({'error': 'Укажите выданную сумму за ремонт', 'field': 'repair_amount'}, ensure_ascii=False)}
            if not parts_name:
                cur.close(); conn.close()
                return {'statusCode': 400, 'headers': HEADERS, 'body': json.dumps({'error': 'Укажите купленную запчасть', 'field': 'parts_name'}, ensure_ascii=False)}

        # Рассчитываем доход мастера = 50% от (repair_amount - purchase_amount)
        master_income_val = 'NULL'
        if repair_amount is not None and purchase_amount is not None:
            profit_raw = int(repair_amount) - int(purchase_amount)
            master_income_calc = max(0, round(profit_raw * 0.5))
            master_income_val = str(master_income_calc)

        # Строим SET-часть динамически
        sets = []
        if new_status:
            sets.append(f"status = '{new_status}'")
            sets.append("status_updated_at = NOW()")
            if new_status in ('ready', 'done'):
                sets.append("completed_at = NOW()")
        if admin_note is not None:
            note_escaped = str(admin_note).replace("'", "''")
            sets.append(f"admin_note = '{note_escaped}'")
        if purchase_amount is not None:
            sets.append(f"purchase_amount = {int(purchase_amount)}")
        if repair_amount is not None:
            sets.append(f"repair_amount = {int(repair_amount)}")
        if parts_name is not None:
            parts_escaped = str(parts_name).replace("'", "''")
            sets.append(f"parts_name = '{parts_escaped}'")
        if master_income_val != 'NULL':
            sets.append(f"master_income = {master_income_val}")
        # Базовые поля заявки
        if upd_name is not None:
            sets.append(f"name = '{str(upd_name).replace(chr(39), chr(39)*2)}'")
        if upd_phone is not None:
            sets.append(f"phone = '{str(upd_phone).replace(chr(39), chr(39)*2)}'")
        if upd_model is not None:
            if upd_model:
                sets.append(f"model = '{str(upd_model).replace(chr(39), chr(39)*2)}'")
            else:
                sets.append("model = NULL")
        if upd_repair_type is not None:
            if upd_repair_type:
                sets.append(f"repair_type = '{str(upd_repair_type).replace(chr(39), chr(39)*2)}'")
            else:
                sets.append("repair_type = NULL")
        if upd_price is not None:
            sets.append(f"price = {int(upd_price)}" if upd_price else "price = NULL")
        if upd_comment is not None:
            if upd_comment:
                sets.append(f"comment = '{str(upd_comment).replace(chr(39), chr(39)*2)}'")
            else:
                sets.append("comment = NULL")

        if not sets:
            cur.close(); conn.close()
            return {'statusCode': 400, 'headers': HEADERS, 'body': json.dumps({'error': 'Нет данных для обновления'}, ensure_ascii=False)}

        set_clause = ', '.join(sets)
        cur.execute(f"UPDATE {SCHEMA}.repair_orders SET {set_clause} WHERE id = {order_id} RETURNING id, name, phone, model, repair_type, repair_amount")
        row = cur.fetchone()
        conn.commit()

        if not row:
            cur.close(); conn.close()
            return {'statusCode': 404, 'headers': HEADERS, 'body': json.dumps({'error': 'Заявка не найдена'}, ensure_ascii=False)}

        # Отправляем Telegram уведомление при смене статуса
        if new_status:
            client_name = row[1] or ''
            client_phone = row[2] or ''
            device_model = row[3] or ''
            repair_t = row[4] or ''
            r_amount = row[5]
            status_label = STATUS_LABELS.get(new_status, new_status)
            tg_msg = (
                f"🔔 *Ремонт #{order_id} — Статус изменён*\n\n"
                f"📱 *Устройство:* {device_model or '—'}\n"
                f"🔧 *Тип ремонта:* {repair_t or '—'}\n"
                f"👤 *Клиент:* {client_name}\n"
                f"📞 *Телефон:* {client_phone}\n"
                f"📌 *Статус:* {status_label}"
                + (f"\n💰 *Стоимость ремонта:* {r_amount} ₽" if r_amount else "")
            )
            token = os.environ['TELEGRAM_BOT_TOKEN']
            main_chat_id = os.environ['TELEGRAM_CHAT_ID']
            send_tg_all(token, main_chat_id, conn, tg_msg)

            # SMS клиенту при смене статуса
            dev = device_model or 'устройство'
            default_templates = {
                'in_progress': 'Скупка24: {device} в ремонте. Готово — сообщим. Skypka24.com',
                'waiting_parts': 'Скупка24: {device} — ждём запчасть. Готово — сообщим. Skypka24.com',
                'ready': 'Скупка24: {device} готов! Стоимость: {amount} руб. Ждём вас. Skypka24.com',
                'done': 'Скупка24: {device} выдан. Спасибо за обращение! Skypka24.com',
                'cancelled': 'Скупка24: {device} — ремонт отменён. Позвоните нам. Skypka24.com',
            }
            if client_phone and new_status in default_templates:
                cur2 = conn.cursor()
                cur2.execute(f"SELECT value FROM {SCHEMA}.settings WHERE key = 'sms_tpl_{new_status}'")
                row2 = cur2.fetchone()
                cur2.close()
                tpl = (row2[0] if row2 and row2[0] else default_templates[new_status])
                sms_text = tpl.replace('{device}', dev).replace('{amount}', str(r_amount or ''))
                send_sms(client_phone, sms_text)

        cur.close(); conn.close()
        return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'ok': True, 'master_income': int(master_income_val) if master_income_val != 'NULL' else None}, ensure_ascii=False)}

    return {'statusCode': 405, 'headers': HEADERS, 'body': json.dumps({'error': 'Method not allowed'}, ensure_ascii=False)}