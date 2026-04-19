import io
import json
import os
import smtplib
import requests
import psycopg2
from email.mime.text import MIMEText

HEADERS = {'Access-Control-Allow-Origin': '*'}
SCHEMA = 't_p31606708_tech_buying_service'
AD_FOOTER = "\n\n📲 Присоединяйтесь к нам: https://t.me/ProService40"

STATUS_LABEL = {
    'in_progress':   'В работе',
    'waiting_parts': 'Ждём запчасть',
    'ready':         'Готово ✓',
    'done':          'Выдано',
    'cancelled':     'Отменено',
}

STATUS_MSG = {
    'in_progress':   'В работе 🔧 Ваш телефон принят и сейчас в ремонте. Сообщим, как только будет готово.',
    'waiting_parts': 'Ждём запчасть ⏳ Запчасть заказана, ожидаем поставку. Сразу приступим к ремонту.',
    'ready':         'Готово ✓ 🎉 Ваш телефон готов! Можно забирать в любое время.',
    'done':          'Выдано 👍 Спасибо за обращение! Рады видеть вас снова.',
    'cancelled':     'Отменено ❌ К сожалению, ремонт отменён. Свяжитесь с нами для уточнения деталей.',
}


def auth_staff(event: dict) -> bool:
    hdrs = event.get('headers') or {}
    admin_token = os.environ.get('ADMIN_TOKEN', '')
    emp_raw = os.environ.get('EMPLOYEE_TOKENS', '')
    valid = [t.strip() for t in emp_raw.split(',') if t.strip()]
    valid.append(admin_token)
    return hdrs.get('X-Admin-Token') in valid or hdrs.get('X-Employee-Token') in valid


def send_tg(token: str, chat_id, text: str, parse_mode: str = 'Markdown'):
    try:
        requests.post(
            f'https://api.telegram.org/bot{token}/sendMessage',
            json={'chat_id': chat_id, 'text': text, 'parse_mode': parse_mode},
            timeout=10,
        )
    except Exception:
        pass


def build_act_docx(order_id, name, phone, model, repair_type, price_str, comment) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import datetime

    doc = Document()

    # Поля страницы
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    def add_para(text='', bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
        return p

    def add_field(label, value):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p.add_run(label)
        r1.bold = True
        r1.font.size = Pt(12)
        r1.font.name = 'Times New Roman'
        r2 = p.add_run(value)
        r2.font.size = Pt(12)
        r2.font.name = 'Times New Roman'
        return p

    now = datetime.datetime.now().strftime('%d.%m.%Y %H:%M')

    add_para('АКТ ПРИЁМКИ НА РЕМОНТ ТЕХНИЧЕСКОГО ОБОРУДОВАНИЯ', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(f'№ {order_id}', size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(f'г. Калуга, {now}', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    add_field('Исполнитель: ', 'ИП Мамедов Адиль Мирза Оглы, ИНН 402810962699, г. Калуга, ул. Кирова, 21а')
    add_field('Заказчик (клиент): ', name)
    add_field('Телефон: ', phone)
    if model:
        add_field('Устройство: ', model)
    if repair_type:
        add_field('Вид работ: ', repair_type)
    add_field('Предварительная стоимость: ', price_str)
    if comment:
        add_field('Описание неисправности: ', comment)

    doc.add_paragraph()
    add_para('УСЛОВИЯ РЕМОНТА — КЛИЕНТ ОЗНАКОМЛЕН И СОГЛАСЕН:', bold=True, size=13)
    doc.add_paragraph()

    risks = [
        'После контакта с водой аппарат может полностью выйти из строя при любом виде ремонта (чистка, прогрев, пайка). Мастерская не несёт ответственности и не обязана восстанавливать устройство.',
        'Компонентная пайка сопряжена с риском безвозвратного повреждения платы. Если телефон перестал включаться в процессе ремонта — работа оплачивается в полном объёме.',
        'При снятии дисплея возможно его повреждение: появление полос, артефактов, отказ включения. Замена повреждённого дисплея производится исключительно за счёт клиента.',
        'Все пользовательские данные (фотографии, контакты, переписка и пр.) могут быть безвозвратно утеряны. Мастерская не занимается восстановлением данных. Подписывая акт, клиент подтверждает, что создал резервную копию либо согласен на потерю данных.',
        'Гарантия на результат ремонта не предоставляется. В худшем случае устройство будет возвращено в нерабочем состоянии; клиент обязуется оплатить стоимость диагностики и всех выполненных работ.',
    ]
    for i, risk in enumerate(risks, 1):
        p = doc.add_paragraph(style='List Number')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'{i}. {risk}')
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    doc.add_paragraph()
    add_para('Подписывая настоящий акт, клиент подтверждает, что ознакомлен со всеми условиями и добровольно соглашается на проведение ремонта.', size=11)

    doc.add_paragraph()
    doc.add_paragraph()

    # Таблица подписей
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            for border in ['top', 'bottom', 'left', 'right']:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                b = OxmlElement(f'w:{border}')
                b.set(qn('w:val'), 'nil')
                tcBorders.append(b)
                tcPr.append(tcBorders)

    # Линии подписей
    for col_idx, label in [(0, 'Мастер (подпись / ФИО)'), (2, 'Клиент (подпись / ФИО)')]:
        cell = table.cell(0, col_idx)
        p = cell.paragraphs[0]
        run = p.add_run('_' * 30)
        run.font.size = Pt(12)
        cell2 = table.cell(1, col_idx)
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(label)
        run2.font.size = Pt(10)
        run2.font.name = 'Times New Roman'

    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_footer.add_run('ИНН: 402810962699  ·  ОГРНИП: 307402814200032  ·  Р/с: 40802810422270001866\nКАЛУЖСКОЕ ОТДЕЛЕНИЕ N8608 ПАО СБЕРБАНК  ·  БИК: 042908612  ·  К/с: 30101810100000000612')
    r.font.size = Pt(9)
    r.font.name = 'Times New Roman'

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def send_tg_document(token: str, chat_id, doc_bytes: bytes, filename: str, caption: str = ''):
    try:
        requests.post(
            f'https://api.telegram.org/bot{token}/sendDocument',
            data={'chat_id': chat_id, 'caption': caption},
            files={'document': (filename, doc_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')},
            timeout=30,
        )
    except Exception:
        pass


def send_sms(phone: str, text: str):
    api_id = os.environ.get('SMSRU_API_ID', '')
    if not api_id:
        return
    try:
        requests.get('https://sms.ru/sms/send',
            params={'api_id': api_id, 'to': phone, 'msg': text, 'json': 1}, timeout=10)
    except Exception:
        pass


def send_email(to: str, subject: str, body: str):
    login = os.environ.get('YANDEX_SMTP_LOGIN', '')
    password = os.environ.get('YANDEX_SMTP_PASSWORD', '')
    if not login or not password:
        return
    try:
        msg = MIMEText(body, 'plain', 'utf-8')
        msg['Subject'] = subject
        msg['From'] = login
        msg['To'] = to
        with smtplib.SMTP_SSL('smtp.yandex.ru', 465) as server:
            server.login(login, password)
            server.sendmail(login, [to], msg.as_string())
    except Exception:
        pass


def save_phone_map(cur, phone: str, chat_id: int, username: str, first_name: str):
    clean = '+' + ''.join(c for c in phone if c.isdigit())
    uname = (username or '').replace("'", "''")
    fname = (first_name or '').replace("'", "''")
    cur.execute(f"""
        INSERT INTO {SCHEMA}.tg_phone_map (phone, chat_id, username, first_name, updated_at)
        VALUES ('{clean}', {chat_id}, '{uname}', '{fname}', NOW())
        ON CONFLICT (phone) DO UPDATE SET
            chat_id=EXCLUDED.chat_id, username=EXCLUDED.username,
            first_name=EXCLUDED.first_name, updated_at=NOW()
    """)
    suffix = clean[-10:]
    cur.execute(f"""
        UPDATE {SCHEMA}.repair_orders
        SET client_tg_chat_id = {chat_id}
        WHERE RIGHT(REGEXP_REPLACE(phone, '[^0-9]', '', 'g'), 10) = '{suffix}'
          AND client_tg_chat_id IS NULL
    """)


FUNC_URL = "https://functions.poehali.dev/8d0ee3bd-41eb-44fe-9d30-aab6ddc2042d"


def handler(event: dict, context) -> dict:
    """
    POST action=new_order      — создать заявку, уведомить команду
    POST action=notify         — отправить клиенту статус через бота (staff only)
    POST action=tg_webhook     — webhook Telegram бота @Skypkaklgbot
    GET  ?setup_webhook=1      — зарегистрировать webhook в Telegram
    """
    if event.get('httpMethod') == 'OPTIONS':
        return {'statusCode': 200, 'headers': {**HEADERS,
            'Access-Control-Allow-Methods': 'GET, POST, OPTIONS',
            'Access-Control-Allow-Headers': 'Content-Type, X-Admin-Token, X-Employee-Token'}, 'body': ''}

    # Регистрация webhook
    params = event.get('queryStringParameters') or {}
    if params.get('setup_webhook'):
        token = os.environ.get('TELEGRAM_BOT_TOKEN', '')
        webhook_url = f"{FUNC_URL}?action=tg_webhook"
        resp = requests.post(
            f'https://api.telegram.org/bot{token}/setWebhook',
            json={'url': webhook_url, 'allowed_updates': ['message']},
            timeout=10,
        )
        return {'statusCode': 200, 'headers': HEADERS,
                'body': json.dumps(resp.json(), ensure_ascii=False)}

    raw = event.get('body') or '{}'
    body = json.loads(raw) if isinstance(raw, str) else (raw or {})
    action = body.get('action', '') or params.get('action', 'new_order')
    token = os.environ.get('TELEGRAM_BOT_TOKEN', '')

    # ── Telegram Webhook от бота ──────────────────────────────────────────────
    if action == 'tg_webhook':
        message = body.get('message') or {}
        chat_id = message.get('chat', {}).get('id')
        if not chat_id or not token:
            return {'statusCode': 200, 'headers': HEADERS, 'body': '{"ok":true}'}

        from_user = message.get('from') or {}
        username = from_user.get('username', '')
        first_name = from_user.get('first_name', '')
        text = (message.get('text') or '').strip()
        contact = message.get('contact')

        conn = psycopg2.connect(os.environ['DATABASE_URL'])
        cur = conn.cursor()

        if contact and contact.get('phone_number'):
            save_phone_map(cur, contact['phone_number'], chat_id, username, first_name)
            conn.commit()
            cur.close(); conn.close()
            send_tg(token, chat_id,
                f"✅ Телефон {contact['phone_number']} привязан.\n"
                "Теперь вы будете получать уведомления о статусе ремонта.", parse_mode='')
            return {'statusCode': 200, 'headers': HEADERS, 'body': '{"ok":true}'}

        if text.startswith('/start'):
            cur.close(); conn.close()
            requests.post(f'https://api.telegram.org/bot{token}/sendMessage', json={
                'chat_id': chat_id,
                'text': (
                    f"Привет, {first_name}! 👋 Я бот Скупка24.\n\n"
                    "Здесь вы будете получать уведомления о статусе ремонта вашего телефона.\n\n"
                    "Нажмите кнопку ниже, чтобы привязать номер 👇"
                ),
                'reply_markup': {
                    'keyboard': [[{'text': '📱 Поделиться номером телефона', 'request_contact': True}]],
                    'resize_keyboard': True, 'one_time_keyboard': True,
                }
            }, timeout=10)
            return {'statusCode': 200, 'headers': HEADERS, 'body': '{"ok":true}'}

        if text.startswith('/status'):
            cur.execute(f"""
                SELECT id, repair_type, status, price FROM {SCHEMA}.repair_orders
                WHERE client_tg_chat_id = {chat_id} ORDER BY id DESC LIMIT 1
            """)
            row = cur.fetchone()
            cur.close(); conn.close()
            if row:
                oid, rtype, status, price = row
                slabels = {'new': 'Новая', 'in_progress': 'В работе 🔧',
                    'waiting_parts': 'Ждём запчасть ⏳', 'ready': 'Готово ✓ 🎉',
                    'done': 'Выдано 👍', 'cancelled': 'Отменено ❌'}
                price_str = f"{int(price):,} ₽".replace(',', ' ') if price else '—'
                send_tg(token, chat_id,
                    f"📋 Заявка #{oid}\nТип: {rtype or '—'}\nСтоимость: {price_str}\n"
                    f"Статус: {slabels.get(status, status)}", parse_mode='')
            else:
                cur.close(); conn.close()
                send_tg(token, chat_id,
                    "Заявок не найдено. Сначала привяжите номер телефона командой /start", parse_mode='')
            return {'statusCode': 200, 'headers': HEADERS, 'body': '{"ok":true}'}

        cur.close(); conn.close()
        send_tg(token, chat_id,
            "Команды:\n/start — привязать номер\n/status — статус ремонта\n\n"
            f"📲 https://t.me/ProService40", parse_mode='')
        return {'statusCode': 200, 'headers': HEADERS, 'body': '{"ok":true}'}

    # ── SMS клиенту о статусе ────────────────────────────────────────────────
    if action == 'notify_sms':
        if not auth_staff(event):
            return {'statusCode': 401, 'headers': HEADERS,
                    'body': json.dumps({'error': 'Unauthorized'}, ensure_ascii=False)}
        order_id = int(body.get('order_id', 0))
        status_key = str(body.get('status_key', '')).strip()
        if not order_id or status_key not in STATUS_MSG:
            return {'statusCode': 400, 'headers': HEADERS,
                    'body': json.dumps({'error': 'order_id и status_key обязательны'}, ensure_ascii=False)}
        conn = psycopg2.connect(os.environ['DATABASE_URL'])
        cur = conn.cursor()
        cur.execute(f"SELECT id, name, phone, repair_type, repair_amount FROM {SCHEMA}.repair_orders WHERE id = {order_id}")
        row = cur.fetchone()
        cur.close(); conn.close()
        if not row:
            return {'statusCode': 404, 'headers': HEADERS, 'body': json.dumps({'error': 'Заявка не найдена'}, ensure_ascii=False)}
        _, name, phone, repair_type, repair_amount = row
        if not phone or not phone.startswith('+7'):
            return {'statusCode': 400, 'headers': HEADERS, 'body': json.dumps({'error': f'Телефон клиента не указан или не в формате +7: {phone}'}, ensure_ascii=False)}
        dev = repair_type or 'устройство'
        sms_templates = {
            'in_progress':   f"Скупка24: {dev} в ремонте. Готово — сообщим. Skypka24.com",
            'waiting_parts': f"Скупка24: {dev} — ждём запчасть. Готово — сообщим. Skypka24.com",
            'ready':         f"Скупка24: {dev} готов! Стоимость: {int(repair_amount) if repair_amount else '?'} руб. Ждём вас. Skypka24.com",
            'done':          f"Скупка24: {dev} выдан. Спасибо за обращение! Skypka24.com",
            'cancelled':     f"Скупка24: {dev} — ремонт отменён. Позвоните нам. Skypka24.com",
        }
        sms_text = sms_templates.get(status_key, '')
        api_id = os.environ.get('SMSRU_API_ID', '')
        if not api_id:
            return {'statusCode': 500, 'headers': HEADERS, 'body': json.dumps({'error': 'SMSRU_API_ID не задан'}, ensure_ascii=False)}
        resp = requests.get('https://sms.ru/sms/send',
            params={'api_id': api_id, 'to': phone, 'msg': sms_text, 'json': 1, 'from': 'IPMamedov'}, timeout=10)
        d = resp.json() if resp.status_code == 200 else {}
        if d.get('status') == 'OK':
            return {'statusCode': 200, 'headers': HEADERS, 'body': json.dumps({'ok': True, 'sent_to': name, 'phone': phone}, ensure_ascii=False)}
        else:
            return {'statusCode': 500, 'headers': HEADERS, 'body': json.dumps({'error': d.get('status_text', 'Ошибка SMS')}, ensure_ascii=False)}

    # ── Уведомить клиента о статусе ──────────────────────────────────────────
    if action == 'notify':
        if not auth_staff(event):
            return {'statusCode': 401, 'headers': HEADERS,
                    'body': json.dumps({'error': 'Unauthorized'}, ensure_ascii=False)}

        order_id = int(body.get('order_id', 0))
        status_key = str(body.get('status_key', '')).strip()

        if not order_id or status_key not in STATUS_MSG:
            return {'statusCode': 400, 'headers': HEADERS,
                    'body': json.dumps({'error': 'order_id и status_key обязательны'}, ensure_ascii=False)}
        if not token:
            return {'statusCode': 500, 'headers': HEADERS,
                    'body': json.dumps({'error': 'TELEGRAM_BOT_TOKEN не задан'}, ensure_ascii=False)}

        conn = psycopg2.connect(os.environ['DATABASE_URL'])
        cur = conn.cursor()
        cur.execute(f"""
            SELECT id, name, phone, repair_type, repair_amount, client_tg_chat_id
            FROM {SCHEMA}.repair_orders WHERE id = {order_id}
        """)
        row = cur.fetchone()
        if not row:
            cur.close(); conn.close()
            return {'statusCode': 404, 'headers': HEADERS,
                    'body': json.dumps({'error': 'Заявка не найдена'}, ensure_ascii=False)}

        _, name, phone, repair_type, repair_amount, chat_id = row

        if not chat_id and phone:
            suffix = ''.join(c for c in phone if c.isdigit())[-10:]
            cur.execute(f"""
                SELECT chat_id FROM {SCHEMA}.tg_phone_map
                WHERE RIGHT(REGEXP_REPLACE(phone, '[^0-9]', '', 'g'), 10) = '{suffix}' LIMIT 1
            """)
            prow = cur.fetchone()
            if prow:
                chat_id = prow[0]

        cur.close(); conn.close()

        if not chat_id:
            return {'statusCode': 404, 'headers': HEADERS,
                    'body': json.dumps({
                        'error': f'Клиент не писал боту @Skypkaklgbot. Попросите клиента написать /start боту.',
                        'phone': phone,
                    }, ensure_ascii=False)}

        msg = STATUS_MSG[status_key]
        amount_str = f"\nСтоимость ремонта: {int(repair_amount):,} ₽".replace(',', ' ') if repair_amount else ""
        text = f"🔧 *Скупка24 — ремонт #{order_id}*\n\n{msg}{amount_str}{AD_FOOTER}"

        resp = requests.post(
            f'https://api.telegram.org/bot{token}/sendMessage',
            json={'chat_id': chat_id, 'text': text, 'parse_mode': 'Markdown'},
            timeout=10,
        )
        tg_data = resp.json()
        if tg_data.get('ok'):
            return {'statusCode': 200, 'headers': HEADERS,
                    'body': json.dumps({'ok': True, 'sent_to': name, 'status': STATUS_LABEL[status_key]}, ensure_ascii=False)}
        else:
            return {'statusCode': 500, 'headers': HEADERS,
                    'body': json.dumps({'error': tg_data.get('description', 'Ошибка Telegram')}, ensure_ascii=False)}

    # ── Создать новую заявку ─────────────────────────────────────────────────
    name = body.get('name', '').strip()
    phone = body.get('phone', '').strip()
    model = body.get('model', '').strip()
    repair_type = body.get('repair_type', '').strip()
    price = body.get('price')
    comment = body.get('comment', '').strip()

    if not name or not phone:
        return {'statusCode': 400, 'headers': HEADERS,
                'body': json.dumps({'error': 'Имя и телефон обязательны'}, ensure_ascii=False)}

    conn = psycopg2.connect(os.environ['DATABASE_URL'])
    cur = conn.cursor()

    # Ищем chat_id по телефону
    suffix = ''.join(c for c in phone if c.isdigit())[-10:]
    cur.execute(f"""
        SELECT chat_id FROM {SCHEMA}.tg_phone_map
        WHERE RIGHT(REGEXP_REPLACE(phone, '[^0-9]', '', 'g'), 10) = '{suffix}' LIMIT 1
    """)
    prow = cur.fetchone()
    client_chat_id = prow[0] if prow else None

    cur.execute(
        f"INSERT INTO {SCHEMA}.repair_orders (name, phone, model, repair_type, price, comment, client_tg_chat_id) "
        "VALUES (%s, %s, %s, %s, %s, %s, %s) RETURNING id",
        (name, phone, model or None, repair_type or None, price, comment or None, client_chat_id)
    )
    order_id = cur.fetchone()[0]
    conn.commit()
    cur.close()
    conn.close()

    price_str = f'{int(price):,} ₽'.replace(',', ' ') if price else 'не определена'
    act_text = (
        f"📋 *АКТ ПРИЁМКИ №{order_id}* | Скупка24, Калуга\n\n"
        f"*Клиент:* {name} | *Тел:* {phone}\n"
        f"*Устройство:* {model or '—'} | *Работы:* {repair_type or '—'}\n"
        f"*Стоимость:* {price_str}"
        + (f" | *Коммент:* {comment}" if comment else "")
        + "\n\n*⚠️ Клиент ознакомлен и согласен:*\n"
        "1. После воды аппарат может умереть при любом ремонте — мастерская не обязана оживлять.\n"
        "2. Компонентная пайка — риск гибели платы; если умер в процессе — работа оплачивается.\n"
        "3. Снятие дисплея — риск полос/артефактов; замена за счёт клиента.\n"
        "4. Данные (фото, контакты) могут быть утеряны безвозвратно.\n"
        "5. Гарантии нет; в худшем случае оплачивается диагностика и выполненные работы.\n\n"
        f"*Подпись клиента:* ______________"
    )
    tg_text = (
        f"🔧 *Заявка #{order_id} на ремонт — Скупка24*\n\n"
        f"👤 *Имя:* {name}\n📞 *Телефон:* {phone}\n"
        f"📱 *Модель:* {model or '—'}\n🛠 *Тип ремонта:* {repair_type or '—'}\n"
        f"💰 *Стоимость:* {price_str}"
        + (f"\n📝 *Комментарий:* {comment}" if comment else "")
        + f"\n\n🔑 ID заявки: `{order_id}`"
        + (f"\n✅ Клиент в TG" if client_chat_id else "")
    )

    if token:
        conn2 = psycopg2.connect(os.environ['DATABASE_URL'])
        cur2 = conn2.cursor()
        cur2.execute(
            f"SELECT telegram_chat_id FROM {SCHEMA}.notification_recipients WHERE is_active=true AND notify_repair=true"
        )
        chat_ids = [r[0] for r in cur2.fetchall()]
        cur2.close(); conn2.close()
        default_chat = os.environ.get('TELEGRAM_CHAT_ID', '')
        if default_chat and default_chat not in chat_ids:
            chat_ids.insert(0, default_chat)
        docx_bytes = build_act_docx(order_id, name, phone, model, repair_type, price_str, comment)
        filename = f'Акт_приёмки_{order_id}_{name.replace(" ", "_")}.docx'
        for cid in chat_ids:
            send_tg(token, cid, tg_text)
            send_tg_document(token, cid, docx_bytes, filename, caption=f'📋 Акт приёмки №{order_id}')

    if token and client_chat_id:
        client_msg = (
            f"✅ *Заявка #{order_id} принята!*\n\n"
            f"Здравствуйте, {name}!\n"
            f"📱 *Устройство:* {model or '—'}\n"
            f"🛠 *Работы:* {repair_type or '—'}\n"
            f"💰 *Стоимость:* {price_str}\n\n"
            f"Скоро перезвоним. Статус — командой /status{AD_FOOTER}"
        )
        send_tg(token, client_chat_id, client_msg)

    send_sms('+79929990333', f'Заявка #{order_id} на ремонт. {name}, {phone}. {repair_type or model or ""}. Скупка24')
    send_email('lekermanya@yandex.ru', f'Заявка #{order_id} на ремонт — Скупка24',
        f"Заявка #{order_id}\nИмя: {name}\nТелефон: {phone}\nМодель: {model}\nТип: {repair_type}\nСтоимость: {price_str}")

    return {'statusCode': 200, 'headers': HEADERS,
            'body': json.dumps({'ok': True, 'order_id': order_id}, ensure_ascii=False)}